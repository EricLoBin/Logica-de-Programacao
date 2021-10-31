'''
- Receber os seguintes dados (pode ser via console - input()):
    - Nome da Disciplina
    - Carga Horária
    - Assunto da disciplina-1
    - Assunto da disciplina-2
    - Assunto da disciplina-3
    - Assunto da disciplina-4

- Criar um documento (DOCX) com os dados acima apresentando o Layout:
    Disciplina: XXXXXXXXXXX
    Com carga horária de 99 horas.
    Conteúdo da disciplina:
    
    +--------+-----------------------------------+
    | Número |              Assunto              |
    +--------+-----------------------------------+
    |      1 | XXXXXXXXX                         |
    |      2 | XXXXXXXXXXXXXXX                   |
    |      3 | XXXXXXXXXXXXXXXXXXXXXXXX          |
    |      4 | XXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXX |
    +--------+-----------------------------------+

'''

import docx

NomeDaDiciplina = input("Digite o Nome da Disciplina: ")
CargaHoraria = input("Digite o Carga Horária: ")

doc = docx.Document()

doc.add_paragraph("Diciplina: ").add_run(NomeDaDiciplina).bold = True
p = doc.add_paragraph("Com carga horária de ")
p.add_run(CargaHoraria).bold = True
p.add_run(' horas.')

tab = doc.add_table(rows=1, cols=2)
tab.style="Colorful Grid Accent 6"
cels = tab.rows[0].cells
cels[0].text = 'Número'
cels[0].width = 5
cels[1].text = "Assunto"

for i in range(1, 5):
    assunto = input("Qual é o assunto número-{}? ".format(i))
    dados = tab.add_row().cells
    dados[0].text = str(i)
    dados[0].width = 5
    dados[1].text = assunto

doc.save("diciplina.docx")
